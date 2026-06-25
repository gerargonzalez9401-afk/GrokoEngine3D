from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Docs/GrokoEngine_Atajos_Teclado.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="222222"):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, title, rows):
    h = doc.add_paragraph()
    h.style = "Heading 2"
    h.add_run(title)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(1.65), Inches(2.05), Inches(2.75)]

    headers = ["Atajo", "Accion", "Notas"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        set_cell_shading(cell, "E8EEF5")
        set_cell_text(cell, text, bold=True, color="1F4D78")

    for shortcut, action, notes in rows:
        cells = table.add_row().cells
        for i, text in enumerate([shortcut, action, notes]):
            cells[i].width = widths[i]
            set_cell_text(cells[i], text)

    doc.add_paragraph()


def add_bullets(doc, title, items):
    h = doc.add_paragraph()
    h.style = "Heading 2"
    h.add_run(title)
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("GrokoEngine - Atajos de teclado y controles")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(19)
    r.font.color.rgb = RGBColor.from_string("2E74B5")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = subtitle.add_run("Guia rapida del editor")
    s.font.name = "Calibri"
    s.font.size = Pt(10.5)
    s.font.color.rgb = RGBColor.from_string("666666")

    add_table(
        doc,
        "Teclado",
        [
            ("Ctrl + Z", "Deshacer", "Restaura el estado anterior de la escena."),
            ("Ctrl + Y", "Rehacer", "Vuelve al estado siguiente despues de deshacer."),
            ("Ctrl + D", "Duplicar seleccion", "Duplica objetos seleccionados con hijos y componentes."),
            ("Delete / Supr", "Borrar seleccion", "Elimina el objeto o seleccion actual."),
            ("F", "Enfocar objeto", "Mueve la camara hacia el objeto seleccionado."),
            ("W A S D", "Mover camara", "Funciona mientras mantienes click derecho en el viewport."),
            ("Space", "Input de scripts", "Disponible para scripts con Input.GetKeyDown(KeyCode.Space)."),
        ],
    )

    add_table(
        doc,
        "Mouse en viewport",
        [
            ("Click izquierdo", "Seleccionar objeto", "Usa raycast contra el volumen del objeto."),
            ("Shift + click izquierdo", "Multiseleccion", "Agrega o quita objetos de la seleccion."),
            ("Arrastrar en vacio", "Marco de seleccion", "Selecciona objetos dentro del rectangulo."),
            ("Click derecho + mover mouse", "Mirar con camara", "Rota la vista de la camara del editor."),
            ("Click derecho + W A S D", "Navegar escena", "Movimiento libre por el viewport."),
            ("Arrastrar eje de gizmo", "Mover, rotar o escalar", "Respeta Snap si Snap esta activo."),
        ],
    )

    add_table(
        doc,
        "Toolbar y modos",
        [
            ("Mover", "Gizmo de posicion", "Boton de herramienta en la barra superior."),
            ("Rotar", "Gizmo de rotacion", "Boton de herramienta en la barra superior."),
            ("Escalar", "Gizmo de escala", "Boton de herramienta en la barra superior."),
            ("Snap: ON/OFF", "Activar snap", "Afecta mover, rotar y escalar con gizmo."),
            ("VSync: ON/OFF", "Limitar FPS por VSync", "Alterna sincronizacion vertical."),
            ("Guardar Proyecto", "Guardar escena", "Guarda en Assets/Scenes/Main.gscene."),
            ("Apply Prefab", "Guardar cambios al prefab", "Disponible si la instancia viene de un prefab."),
        ],
    )

    add_table(
        doc,
        "Prefabs y Project",
        [
            ("Arrastrar objeto a Project", "Crear prefab", "Crea un archivo .prefab desde la jerarquia."),
            ("Doble click en prefab", "Instanciar prefab", "Agrega una instancia a la escena."),
            ("Click en prefab", "Editar prefab en Inspector", "Permite ver y modificar sus componentes."),
            ("Arrastrar prefab a jerarquia", "Instanciar como hijo", "Lo coloca dentro del objeto destino."),
            ("Arrastrar prefab a viewport", "Soltar en escena", "Detecta suelo/colliders y lo coloca encima."),
            ("Arrastrar script a objeto", "Agregar script", "Suelta un .cs sobre un objeto de la jerarquia."),
        ],
    )

    add_bullets(
        doc,
        "Notas rapidas",
        [
            "Los cubos seleccionados muestran wireframe amarillo en el viewport.",
            "El preview azul de prefab indica donde se soltara el objeto.",
            "Undo/Redo trabaja con snapshots de escena y no modifica archivos hasta guardar.",
            "En Play Mode no se deben editar componentes desde el Inspector.",
        ],
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("GrokoEngine - referencia rapida")
    fr.font.name = "Calibri"
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor.from_string("777777")

    doc.save(OUT)


if __name__ == "__main__":
    main()
